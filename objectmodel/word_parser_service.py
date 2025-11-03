#!/usr/bin/env python3
"""
Service d'extraction des tableaux Word pour ATP
Extrait les tableaux avec leur structure complète
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import tempfile
import os
import json

app = Flask(__name__)
CORS(app)

def extract_tables_from_docx(file_path):
    """
    Extrait tous les tableaux d'un fichier Word
    Retourne une structure avec chapitres, sections et test cases
    """
    doc = Document(file_path)
    results = []
    id_counter = 1
    current_chapter = None
    
    # D'abord extraire les titres/chapitres depuis les paragraphes
    for i, para in enumerate(doc.paragraphs):
        line = para.text.strip()
        if not line:
            continue
        
        # Détecter les chapitres principaux (ex: "7 Test Suites")
        chapter_match = None
        if line and line[0].isdigit():
            parts = line.split(None, 1)
            if len(parts) >= 2 and parts[0].isdigit() and len(parts[0]) <= 2:
                chapter_num = parts[0]
                chapter_title = parts[1] if len(parts) > 1 else ""
                chapter_match = (chapter_num, chapter_title)
        
        if chapter_match:
            chapter_num, chapter_title = chapter_match
            current_chapter = {'number': chapter_num, 'title': chapter_title}
            
            results.append({
                '_id': f'chapter_{chapter_num}_{id_counter}',
                'type': 'chapter',
                'number': chapter_num,
                'title': chapter_title,
                'line': i + 1,
                'timestamp': None  # Sera ajouté côté client
            })
            id_counter += 1
            continue
        
        # Détecter les sous-sections (ex: "6.3.1 WriteAttributes")
        if '.' in line and line.split()[0].replace('.', '').isdigit():
            parts = line.split(None, 1)
            if len(parts) >= 1:
                section_num = parts[0]
                section_title = parts[1] if len(parts) > 1 else ""
                
                test_type = 'procedure' if section_num.startswith('6.') else 'section'
                
                results.append({
                    '_id': f'section_{section_num.replace(".", "_")}_{id_counter}',
                    'type': test_type,
                    'number': section_num,
                    'title': section_title,
                    'line': i + 1,
                    'parent': current_chapter['number'] if current_chapter else None,
                    'timestamp': None
                })
                id_counter += 1
    
    # Maintenant extraire les tableaux
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        
        # Extraire les en-têtes (première ligne)
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # Détecter si c'est un tableau de test case (première cellule = "Test Case")
        if headers and headers[0] == 'Test Case' and len(headers) == 2:
            # C'est un tableau de test case en format 2 colonnes
            test_case_id = headers[1]
            
            # Extraire le test ID et le titre
            if 'ACESM-' in test_case_id and '-TC' in test_case_id:
                # Format: "ACESM-UC06-TC01: Remote connect/disconnect..."
                parts = test_case_id.split(':', 1)
                if len(parts) == 2:
                    test_id = parts[0].strip()
                    test_title = parts[1].strip()
                else:
                    test_id = test_case_id
                    test_title = test_case_id
                
                # Extraire toutes les données du tableau
                table_data = {}
                for row in table.rows[1:]:
                    if len(row.cells) >= 2:
                        key = row.cells[0].text.strip()
                        value = row.cells[1].text.strip()
                        if key:  # Ignorer les clés vides
                            table_data[key] = value
                
                # Créer l'entrée de test case
                test_entry = {
                    '_id': f'test_{test_id}_{id_counter}',
                    'type': 'test-case',
                    'testId': test_id,
                    'title': test_title,
                    'line': -1,  # Pas de ligne pour les tableaux
                    'chapter': current_chapter['number'] if current_chapter else None,
                    'timestamp': None,
                    'tableColumns': headers,
                    'tableData': table_data,
                    # Champs détaillés
                    'useCase': table_data.get('Use Case', ''),
                    'scenario': table_data.get('Scenario', ''),
                    'testPurpose': table_data.get('Test purpose', ''),
                    'testStrategy': table_data.get('Test strategy', ''),
                    'aaFilter': table_data.get('AA filter', ''),
                    'prerequisites': table_data.get('Prerequisites', ''),
                    'preamble': table_data.get('Preamble', ''),
                    'testBody': table_data.get('Test body', ''),
                    'postamble': table_data.get('Postamble', ''),
                    'comment': table_data.get('Comment', ''),
                }
                
                results.append(test_entry)
                id_counter += 1
        
        else:
            # Tableau général (multi-colonnes)
            # Créer une entrée pour chaque ligne de données
            for row_idx, row in enumerate(table.rows[1:]):  # Skip header
                cells_text = [cell.text.strip() for cell in row.cells]
                
                # Skip empty rows
                if not any(cells_text):
                    continue
                
                # Vérifier si cette ligne contient un test case (format ACESM-*-TC*)
                has_test_case = any('ACESM-' in cell and '-TC' in cell for cell in cells_text)
                
                if has_test_case:
                    # Trouver le test ID
                    test_id = None
                    for cell in cells_text:
                        if 'ACESM-' in cell and '-TC' in cell:
                            # Extraire le test ID
                            parts = cell.split()
                            for part in parts:
                                if 'ACESM-' in part and '-TC' in part:
                                    test_id = part.rstrip(':,.')
                                    break
                            if test_id:
                                break
                    
                    if test_id:
                        # Créer un dictionnaire des données de la ligne
                        row_data = dict(zip(headers, cells_text))
                        
                        results.append({
                            '_id': f'test_table_{test_id}_{id_counter}',
                            'type': 'test-case',
                            'testId': test_id,
                            'title': ' | '.join(cells_text[:3]),  # Première colonnes comme titre
                            'line': -1,
                            'chapter': current_chapter['number'] if current_chapter else None,
                            'timestamp': None,
                            'tableColumns': headers,
                            'tableData': row_data,
                            'tableRows': [cells_text]
                        })
                        id_counter += 1
    
    return results


@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok'})


@app.route('/parse', methods=['POST'])
def parse_document():
    """
    Endpoint pour parser un document Word
    Accepte un fichier .docx en multipart/form-data
    Retourne la structure JSON des tests extraits
    """
    try:
        # Vérifier qu'un fichier est présent
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.endswith('.docx'):
            return jsonify({'error': 'Only .docx files are supported'}), 400
        
        # Sauvegarder temporairement le fichier
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
        
        try:
            # Parser le document
            results = extract_tables_from_docx(temp_path)
            
            return jsonify({
                'success': True,
                'count': len(results),
                'tests': results
            })
        
        finally:
            # Nettoyer le fichier temporaire
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


if __name__ == '__main__':
    print("Starting ATP Word Parser Service...")
    print("Service running on http://localhost:5000")
    print("Endpoints:")
    print("  - GET  /health  : Health check")
    print("  - POST /parse   : Parse Word document")
    app.run(debug=True, host='0.0.0.0', port=5000)
